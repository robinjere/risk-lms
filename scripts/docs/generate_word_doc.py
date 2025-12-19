"""
Generate Word Document for Risk LMS System Documentation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title Page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title_run = title.add_run('RISK DEPARTMENT')
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('LEARNING MANAGEMENT SYSTEM')
subtitle_run.bold = True
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 51, 102)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

org = doc.add_paragraph()
org_run = org.add_run('Co-operative Bank of Tanzania PLC')
org_run.bold = True
org_run.font.size = Pt(18)
org.alignment = WD_ALIGN_PARAGRAPH.CENTER

dept = doc.add_paragraph()
dept_run = dept.add_run('Risk Management & Compliance Department')
dept_run.font.size = Pt(14)
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Document info table
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Document Title:', 'System Documentation'),
    ('Version:', '1.0'),
    ('Date:', 'December 19, 2025'),
    ('Classification:', 'Internal Use'),
    ('Author:', 'System Development Team'),
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Overview',
    '3. Technology Stack',
    '4. User Roles & Permissions',
    '5. System Architecture',
    '6. Core Modules',
    '7. Authentication System',
    '8. Course Management',
    '9. Video & Interactive Content',
    '10. Quiz & Assessment System',
    '11. Certificate Generation',
    '12. Database Schema',
    '13. Installation & Deployment',
    '14. Administration Guide',
    '15. User Guide',
    '16. Security Considerations',
    '17. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# Section 1: Executive Summary
doc.add_heading('1. Executive Summary', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph('The Risk Department Learning Management System (Risk LMS) is a comprehensive web-based training platform designed specifically for Co-operative Bank of Tanzania PLC. The system enables the Risk Management & Compliance department to deliver mandatory training courses to all bank staff (bankers), track their progress, assess their knowledge through quizzes, and issue verifiable certificates upon successful completion.')

doc.add_heading('1.2 Key Objectives', level=2)
objectives = [
    'Centralized Training Management: Single platform for all risk and compliance training materials',
    'Mandatory Viewing Enforcement: Ensures staff watch 95% of video content without skipping',
    'Knowledge Assessment: Random question pools with 80% pass benchmark',
    'Compliance Tracking: Real-time monitoring of training completion rates',
    'Verifiable Certification: QR-coded certificates for authenticity verification',
    'Active Directory Integration: Seamless authentication with bank domain credentials',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.3 Target Users', level=2)
users = [
    'Head of Risk: Full administrative access to manage courses, users, and view analytics',
    'Risk & Compliance Specialist: Course creation, content management, and user monitoring',
    'Bankers: Training participants who complete courses and earn certificates',
]
for user in users:
    doc.add_paragraph(user, style='List Bullet')

doc.add_page_break()

# Section 2: System Overview
doc.add_heading('2. System Overview', level=1)

doc.add_heading('2.1 System Description', level=2)
doc.add_paragraph('The Risk LMS is built on Django 4.2+, a robust Python web framework, with Microsoft SQL Server as the production database. The system features a modern, responsive interface based on SB Admin 2 Bootstrap template.')

doc.add_heading('2.2 Key Features', level=2)
features_table = doc.add_table(rows=9, cols=2)
features_table.style = 'Table Grid'
features_data = [
    ('Feature', 'Description'),
    ('Multi-format Content', 'Support for video courses, interactive SCORM/Captivate packages, and HTML5 content'),
    ('Progress Tracking', 'Real-time tracking of user progress with mandatory viewing requirements'),
    ('Random Quiz Generation', 'Questions randomly selected from topic-based question pools'),
    ('Certificate Generation', 'Automated PDF certificates with QR codes for verification'),
    ('Active Directory Auth', 'Single Sign-On with bank KCBLTZ domain'),
    ('Role-Based Access', 'Four-tier permission system for different user types'),
    ('Analytics Dashboard', 'Comprehensive reporting on training compliance'),
    ('Deadline Management', 'Configurable course completion time limits'),
]
for i, (feature, desc) in enumerate(features_data):
    features_table.rows[i].cells[0].text = feature
    features_table.rows[i].cells[1].text = desc
    if i == 0:
        for cell in features_table.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_page_break()

# Section 3: Technology Stack
doc.add_heading('3. Technology Stack', level=1)

doc.add_heading('3.1 Backend Technologies', level=2)
backend_table = doc.add_table(rows=9, cols=3)
backend_table.style = 'Table Grid'
backend_data = [
    ('Component', 'Technology', 'Version'),
    ('Framework', 'Django', '4.2.x'),
    ('Language', 'Python', '3.9+'),
    ('Database', 'Microsoft SQL Server', '2016+'),
    ('Database Driver', 'mssql-django, pyodbc', 'Latest'),
    ('Task Queue', 'Celery', '5.3+'),
    ('Message Broker', 'Redis', '5.0+'),
    ('WSGI Server', 'Gunicorn/Waitress', 'Latest'),
    ('Authentication', 'ldap3', '2.9+'),
]
for i, row_data in enumerate(backend_data):
    for j, cell_data in enumerate(row_data):
        backend_table.rows[i].cells[j].text = cell_data
        if i == 0:
            backend_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_heading('3.2 Frontend Technologies', level=2)
frontend_table = doc.add_table(rows=5, cols=2)
frontend_table.style = 'Table Grid'
frontend_data = [
    ('Component', 'Technology'),
    ('CSS Framework', 'Bootstrap 4 (SB Admin 2)'),
    ('JavaScript', 'jQuery 3.x'),
    ('Charts', 'Chart.js'),
    ('Icons', 'Font Awesome 5'),
]
for i, row_data in enumerate(frontend_data):
    for j, cell_data in enumerate(row_data):
        frontend_table.rows[i].cells[j].text = cell_data
        if i == 0:
            frontend_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# Section 4: User Roles & Permissions
doc.add_heading('4. User Roles & Permissions', level=1)

doc.add_heading('4.1 Role Hierarchy', level=2)
doc.add_paragraph('The system implements a four-tier permission hierarchy:')
roles = [
    'ADMIN: Full System Access - Superuser privileges',
    'HEAD OF RISK: Course Management, User Management, Analytics',
    'RISK & COMPLIANCE SPECIALIST: Course Management, Content Upload, User Monitoring',
    'BANKER: View Courses, Take Quizzes, Earn Certificates',
]
for role in roles:
    doc.add_paragraph(role, style='List Number')

doc.add_heading('4.2 Permission Matrix', level=2)
perm_table = doc.add_table(rows=13, cols=5)
perm_table.style = 'Table Grid'
perm_data = [
    ('Permission', 'Admin', 'Head of Risk', 'R&C Specialist', 'Banker'),
    ('Access Admin Panel', 'Yes', 'Yes', 'Yes', 'No'),
    ('Create Courses', 'Yes', 'Yes', 'Yes', 'No'),
    ('Upload Videos', 'Yes', 'Yes', 'Yes', 'No'),
    ('Add Questions', 'Yes', 'Yes', 'Yes', 'No'),
    ('Delete Courses', 'Yes', 'Yes', 'Yes', 'No'),
    ('Manage Users', 'Yes', 'Yes', 'Yes', 'No'),
    ('View Analytics', 'Yes', 'Yes', 'Yes', 'No'),
    ('Export Reports', 'Yes', 'Yes', 'Yes', 'No'),
    ('Enroll in Courses', 'Yes', 'Yes', 'Yes', 'Yes'),
    ('Take Quizzes', 'Yes', 'Yes', 'Yes', 'Yes'),
    ('View Own Progress', 'Yes', 'Yes', 'Yes', 'Yes'),
    ('Download Certificates', 'Yes', 'Yes', 'Yes', 'Yes'),
]
for i, row_data in enumerate(perm_data):
    for j, cell_data in enumerate(row_data):
        perm_table.rows[i].cells[j].text = cell_data
        if i == 0:
            perm_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# Section 5: System Architecture
doc.add_heading('5. System Architecture', level=1)

doc.add_heading('5.1 Application Structure', level=2)
doc.add_paragraph('The application follows Django best practices with modular app structure:')
apps = [
    'risk_lms/ - Main Django project (settings, URLs, WSGI)',
    'accounts/ - User management and authentication',
    'courses/ - Course creation and enrollment',
    'videos/ - Video and interactive content',
    'quizzes/ - Assessment and question bank',
    'certificates/ - Certificate generation',
    'content_management/ - Admin content tools',
    'templates/ - HTML templates',
    'static/ - CSS, JavaScript, images',
    'media/ - Uploaded files',
]
for app in apps:
    doc.add_paragraph(app, style='List Bullet')

doc.add_heading('5.2 Database Tables', level=2)
tables = [
    'users - User accounts and profiles',
    'courses - Course definitions',
    'enrollments - User-course enrollment records',
    'videos - Video content metadata',
    'video_progress - User video watching progress',
    'interactive_courses - SCORM/Captivate packages',
    'interactive_course_progress - Interactive content progress',
    'questions - Question bank',
    'question_options - Answer options',
    'quiz_attempts - Quiz attempt records',
    'quiz_answers - Individual question answers',
    'certificates - Issued certificates',
]
for table in tables:
    doc.add_paragraph(table, style='List Bullet')

doc.add_page_break()

# Section 6: Core Modules
doc.add_heading('6. Core Modules', level=1)

doc.add_heading('6.1 Accounts Module', level=2)
doc.add_paragraph('Purpose: User authentication, authorization, and profile management.')
doc.add_paragraph('Key Components:')
acc_components = [
    'Custom User model with role-based permissions',
    'Active Directory/LDAP authentication backend',
    'Local password authentication fallback',
    'User profile management',
]
for comp in acc_components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_heading('6.2 Courses Module', level=2)
doc.add_paragraph('Purpose: Course creation, management, and enrollment tracking.')
doc.add_paragraph('Key Components:')
course_components = [
    'Course creation and editing',
    'Enrollment management',
    'Progress tracking',
    'Deadline management',
]
for comp in course_components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_heading('6.3 Videos Module', level=2)
doc.add_paragraph('Purpose: Video content and interactive course management.')
doc.add_paragraph('Supported Interactive Content Types:')
content_types = [
    'Adobe Captivate - .cptx packages exported as HTML5',
    'SCORM 1.2/2004 - Standard eLearning packages',
    'Articulate Storyline - Published HTML5 output',
    'HTML5 - Custom HTML5 courses',
]
for ct in content_types:
    doc.add_paragraph(ct, style='List Bullet')

doc.add_heading('6.4 Quizzes Module', level=2)
doc.add_paragraph('Purpose: Question bank management and quiz assessment.')
doc.add_paragraph('Question Types:')
q_types = [
    'Multiple Choice - Single correct answer',
    'True/False - Boolean answer',
    'Multiple Answer - Multiple correct options',
]
for qt in q_types:
    doc.add_paragraph(qt, style='List Bullet')

doc.add_heading('6.5 Certificates Module', level=2)
doc.add_paragraph('Purpose: Certificate generation and verification.')
doc.add_paragraph('Certificate Features:')
cert_features = [
    'PDF certificate with bank branding',
    'QR code for instant verification',
    'Unique certificate number',
    'Online verification URL',
]
for cf in cert_features:
    doc.add_paragraph(cf, style='List Bullet')

doc.add_page_break()

# Section 7: Authentication
doc.add_heading('7. Authentication System', level=1)

doc.add_heading('7.1 Active Directory Integration', level=2)
doc.add_paragraph('The system integrates with Co-operative Bank Active Directory domain for Single Sign-On authentication.')

doc.add_paragraph('Domain Configuration:')
config_items = [
    'Domain: KCBLTZ.CRDBBANKPLC.COM',
    'Primary DNS: 192.168.10.50',
    'Alternate DNS: 192.168.10.10',
    'Base DN: DC=KCBLTZ,DC=CRDBBANKPLC,DC=COM',
    'Port: 389 (LDAP)',
]
for item in config_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.2 Authentication Flow', level=2)
auth_flow = [
    'User enters username and password',
    'System attempts LDAP authentication against AD servers',
    'If successful, user details are retrieved from AD',
    'Local user record is created/updated',
    'Session is established',
    'User is redirected based on role',
]
for i, step in enumerate(auth_flow, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('7.3 Dual Authentication Support', level=2)
doc.add_paragraph('Users can authenticate using:')
doc.add_paragraph('Domain Credentials: Primary method using AD username/password', style='List Bullet')
doc.add_paragraph('Local Credentials: Fallback for testing or when AD is unavailable', style='List Bullet')

doc.add_page_break()

# Section 8: Course Management
doc.add_heading('8. Course Management', level=1)

doc.add_heading('8.1 Creating a Course', level=2)
doc.add_paragraph('Step-by-step process:')
create_steps = [
    'Navigate to Content Dashboard',
    'Click Create New Course button',
    'Enter course details (Title, Description, Passing Score)',
    'Upload thumbnail image (optional)',
    'Set completion time limit (optional)',
    'Save and proceed to content upload',
]
for i, step in enumerate(create_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('8.2 Course Structure', level=2)
doc.add_paragraph('Each course contains:')
structure = [
    'Course Details - Title, Description, Thumbnail',
    'Settings - Passing Score, Time Limit',
    'Videos - Multiple, ordered video content',
    'Interactive Courses - SCORM/Captivate packages',
    'Questions - Random pool by topic',
    'Enrollments - User progress tracking',
]
for item in structure:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.3 Publishing Workflow', level=2)
doc.add_paragraph('Course publishing workflow:')
workflow = [
    'Create course (Draft status)',
    'Upload content (Videos, Interactive)',
    'Add questions (Minimum 5 required)',
    'Review and test',
    'Publish course',
]
for i, step in enumerate(workflow, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# Section 9: Video & Interactive Content
doc.add_heading('9. Video & Interactive Content', level=1)

doc.add_heading('9.1 Supported Video Formats', level=2)
formats = ['MP4', 'WebM', 'AVI', 'MOV']
for fmt in formats:
    doc.add_paragraph(fmt, style='List Bullet')

doc.add_heading('9.2 Progress Tracking', level=2)
doc.add_paragraph('Video Progress:')
video_tracking = [
    'Tracks watched duration',
    'Records last position for resume',
    'Requires 95% completion',
    'No skipping allowed',
]
for item in video_tracking:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Interactive Course Progress:')
interactive_tracking = [
    'Slide-by-slide tracking',
    'SCORM data persistence',
    'Completion percentage calculation',
    'Resume capability',
]
for item in interactive_tracking:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 10: Quiz System
doc.add_heading('10. Quiz & Assessment System', level=1)

doc.add_heading('10.1 Grading System', level=2)
grading = [
    'Pass Mark: 80% (configurable per course)',
    'Scoring: Points-based (default 1 point per question)',
    'Multiple Attempts: Allowed, best score recorded',
    'Time Limit: None (optional feature)',
]
for item in grading:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Quiz Flow', level=2)
quiz_flow = [
    'Prerequisites Check - Must complete 95% of video content',
    'Generate Questions - Random selection from pool',
    'Present Quiz - Display questions one at a time',
    'Answer Questions - User selects answers',
    'Calculate Score - Determine pass/fail',
    'Generate Certificate - If passed, certificate is issued',
]
for i, step in enumerate(quiz_flow, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# Section 11: Certificate Generation
doc.add_heading('11. Certificate Generation', level=1)

doc.add_heading('11.1 Certificate Components', level=2)
doc.add_paragraph('PDF Certificate includes:')
cert_components = [
    'Bank logo and branding',
    'Recipient full name',
    'Course title',
    'Completion date',
    'Final score',
    'Unique certificate number',
    'QR code for verification',
    'Authorized signatures',
]
for item in cert_components:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('11.2 QR Code Verification', level=2)
doc.add_paragraph('The QR code contains comprehensive verification information including certificate number, recipient name, course title, score, and completion date. Scanning the QR code instantly displays all certificate details for verification.')

doc.add_heading('11.3 Verification Methods', level=2)
verification = [
    'QR Code Scan: Instant verification via mobile device',
    'Online Verification: URL-based verification page',
    'Certificate Number: Manual lookup in system',
]
for item in verification:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 12: Database Schema
doc.add_heading('12. Database Schema', level=1)

doc.add_heading('12.1 User Model', level=2)
user_fields = [
    'username - Domain username (sAMAccountName)',
    'email - Email address (unique)',
    'first_name, last_name - Full name',
    'role - User role (admin, head_of_risk, risk_compliance_specialist, banker)',
    'phone - Contact number',
    'department - Bank department',
    'is_staff - Staff status',
    'is_superuser - Superuser status',
]
for field in user_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('12.2 Course Model', level=2)
course_fields = [
    'title - Course title',
    'description - Course description',
    'created_by - Creator (User FK)',
    'is_published - Publication status',
    'passing_score - Minimum pass percentage (default 80)',
    'completion_time_enabled - Time limit enabled',
    'completion_time_limit - Days to complete',
]
for field in course_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_page_break()

# Section 13: Installation
doc.add_heading('13. Installation & Deployment', level=1)

doc.add_heading('13.1 Prerequisites', level=2)
prereqs = [
    'Python 3.9 or higher',
    'Microsoft SQL Server 2016+',
    'ODBC Driver 17 for SQL Server',
    'Git (optional)',
]
for prereq in prereqs:
    doc.add_paragraph(prereq, style='List Bullet')

doc.add_heading('13.2 Installation Steps', level=2)
install_steps = [
    'Extract or clone project to desired location',
    'Create virtual environment: python -m venv venv',
    'Activate environment: .\\venv\\Scripts\\Activate.ps1',
    'Install dependencies: pip install -r requirements.txt',
    'Set environment variables for database connection',
    'Run migrations: python manage.py migrate',
    'Create superuser: python manage.py createsuperuser',
    'Collect static files: python manage.py collectstatic',
    'Run server: python manage.py runserver 0.0.0.0:8000',
]
for i, step in enumerate(install_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('13.3 Environment Variables', level=2)
env_table = doc.add_table(rows=6, cols=3)
env_table.style = 'Table Grid'
env_data = [
    ('Variable', 'Description', 'Default'),
    ('DB_ENGINE', 'Database engine', 'mssql'),
    ('DB_NAME', 'Database name', 'risk_lms'),
    ('DB_HOST', 'Database server', 'localhost\\SQLEXPRESS'),
    ('DEBUG', 'Debug mode', 'True'),
    ('SECRET_KEY', 'Django secret key', '(Change in production)'),
]
for i, row_data in enumerate(env_data):
    for j, cell_data in enumerate(row_data):
        env_table.rows[i].cells[j].text = cell_data
        if i == 0:
            env_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# Section 14: Administration Guide
doc.add_heading('14. Administration Guide', level=1)

doc.add_heading('14.1 User Management', level=2)
doc.add_paragraph('Creating Users:')
user_mgmt = [
    'Users are auto-created on first AD login',
    'Manual creation via Django Admin',
    'Bulk import via management command',
]
for item in user_mgmt:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('14.2 Course Administration', level=2)
doc.add_paragraph('Course Lifecycle:')
lifecycle = [
    'Create course (Draft status)',
    'Upload content (Videos, Interactive)',
    'Add questions (Minimum 5)',
    'Review and test',
    'Publish course',
    'Monitor enrollments',
    'Archive when complete',
]
for i, step in enumerate(lifecycle, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('14.3 Reporting', level=2)
doc.add_paragraph('Available Reports:')
reports = [
    'User Progress Report',
    'Course Completion Report',
    'Quiz Performance Report',
    'Certificate Issuance Report',
    'Compliance Summary Report',
]
for report in reports:
    doc.add_paragraph(report, style='List Bullet')

doc.add_page_break()

# Section 15: User Guide
doc.add_heading('15. User Guide', level=1)

doc.add_heading('15.1 For Bankers (Training Participants)', level=2)

doc.add_paragraph('Logging In:')
login_steps = [
    'Navigate to http://localhost:8000/',
    'Enter your domain username (e.g., JMURO)',
    'Enter your domain password',
    'Click Sign In',
]
for i, step in enumerate(login_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('Taking a Course:')
course_steps = [
    'View available courses on dashboard',
    'Click Enroll on desired course',
    'Watch all videos (95% minimum)',
    'Complete interactive content',
    'Take the quiz (80% pass required)',
    'Download your certificate',
]
for i, step in enumerate(course_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('15.2 For Administrators', level=2)
doc.add_paragraph('Creating Content:')
admin_steps = [
    'Login with admin credentials',
    'Navigate to Content Dashboard',
    'Create new course',
    'Upload videos/interactive content',
    'Add quiz questions',
    'Publish course',
]
for i, step in enumerate(admin_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# Section 16: Security
doc.add_heading('16. Security Considerations', level=1)

doc.add_heading('16.1 Authentication Security', level=2)
auth_security = [
    'LDAP/AD Integration: Enterprise-grade authentication',
    'Password Hashing: Django PBKDF2 algorithm',
    'Session Security: Secure session cookies',
    'CSRF Protection: Built-in CSRF tokens',
]
for item in auth_security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('16.2 Authorization Controls', level=2)
auth_controls = [
    'Role-Based Access: Four-tier permission system',
    'View Decorators: @login_required protection',
    'Permission Checks: can_upload_content() methods',
    'Admin Panel: Staff-only access',
]
for item in auth_controls:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('16.3 Data Protection', level=2)
data_protection = [
    'SQL Injection: Django ORM parameterized queries',
    'XSS Prevention: Template auto-escaping',
    'File Upload: Validated file types',
    'Sensitive Data: Encrypted storage recommended',
]
for item in data_protection:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('16.4 Production Recommendations', level=2)
prod_recs = [
    'Set DEBUG = False',
    'Use secure SECRET_KEY from environment variable',
    'Enable HTTPS with SSL certificate',
    'Set SECURE_SSL_REDIRECT = True',
    'Configure SESSION_COOKIE_SECURE = True',
    'Enable CSRF_COOKIE_SECURE = True',
]
for item in prod_recs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 17: Troubleshooting
doc.add_heading('17. Troubleshooting', level=1)

doc.add_heading('17.1 Common Issues', level=2)

doc.add_paragraph('Issue: Cannot connect to database')
doc.add_paragraph('Solution: Verify SQL Server is running and environment variables are set correctly. Check DB_HOST format includes instance name (e.g., ServerName\\SQLEXPRESS).')

doc.add_paragraph('Issue: LDAP authentication fails')
doc.add_paragraph('Solution: Check network connectivity to AD servers (192.168.10.50, 192.168.10.10). Verify domain credentials are correct. Ensure LDAP port 389 is open.')

doc.add_paragraph('Issue: Video will not play')
doc.add_paragraph('Solution: Check file format (MP4/WebM supported). Verify file uploaded completely. Check browser console for errors.')

doc.add_paragraph('Issue: Certificate generation fails')
doc.add_paragraph('Solution: Ensure reportlab and Pillow are installed. Check media/certificates/ directory permissions. Verify all required user fields are populated.')

doc.add_heading('17.2 Support Contacts', level=2)
doc.add_paragraph('For technical support, contact the IT Department or System Development Team.')

doc.add_page_break()

# Final page
doc.add_paragraph()
doc.add_paragraph()
final = doc.add_paragraph()
final_run = final.add_run('Document Information')
final_run.bold = True
final_run.font.size = Pt(14)
final.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

info_table2 = doc.add_table(rows=6, cols=2)
info_table2.style = 'Table Grid'
final_info = [
    ('Document Title', 'Risk LMS System Documentation'),
    ('Version', '1.0'),
    ('Last Updated', 'December 19, 2025'),
    ('Organization', 'Co-operative Bank of Tanzania PLC'),
    ('Department', 'Risk Management & Compliance'),
    ('Classification', 'Internal Use'),
]
for i, (label, value) in enumerate(final_info):
    info_table2.rows[i].cells[0].text = label
    info_table2.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph()
copyright_p = doc.add_paragraph()
copyright_p.add_run('© 2025 Co-operative Bank of Tanzania PLC. All rights reserved.')
copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save('Risk_LMS_System_Documentation.docx')
print('=' * 60)
print('Word document created successfully!')
print('File: Risk_LMS_System_Documentation.docx')
print('=' * 60)
